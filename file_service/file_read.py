from os.path import join, dirname, exists
from qrcode import QRCode, constants
from .file_path import get_file_path
from docx.shared import Inches, Pt
from openpyxl import load_workbook
from os.path import join, dirname
from datetime import datetime
from docx2pdf import convert
from docx import Document
from os import remove


async def replace_text(paragraph, old_text, new_text, font_size=None, bold=True, underline=False):
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            if font_size:
                run.font.size = Pt(font_size)
            if bold:
                run.font.bold = True
            if underline:
                run.font.underline = True


async def add_qrcode_pcture(paragraph, old_text, new_text, size=1.6):
    for run in paragraph.runs:
        print(run.text)
        if old_text in run.text:
            run.add_picture(join(dirname(__file__), f"ariza_qrcode/{new_text}.png"),
                            width=Inches(size))


async def replace_table_text(table, old_text, new_text, font_size=None, bold=True, size=1.6):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if old_text in run.text:
                        if old_text == "&":
                            run.add_picture(join(dirname(__file__), f"file_qrcode/{new_text}.png"),
                                            width=Inches(size))
                        if old_text != "&":
                            run.text = run.text.replace(old_text, new_text)
                        if font_size:
                            run.font.size = Pt(font_size)
                        if bold:
                            run.font.bold = True


async def process_document(address, name, file_name):
    doc = Document(join(dirname(__file__), str(file_name)))
    for paragraph in doc.paragraphs:
        if "ADDRESS" in paragraph.text:
            await replace_text(paragraph, "ADDRESS", address)
        if "NAME" in paragraph.text:
            await replace_text(paragraph, "NAME", name)
        if "DATEFULL" in paragraph.text:
            await replace_text(paragraph, "DATEFULL", f"{datetime.now().strftime('%d.%m.%Y')}    {name}")
        if "&" in paragraph.text:
            await add_qrcode_pcture(paragraph, "&", new_text=f"{name}", size=1.6)
    doc.save(join(dirname(__file__), f"file_ariza\\{name}.docx"))
    await convert_pdf(name=name, status=True)


async def process_contract(name, faculty, passport, number, address, contract_number, file_name):
    doc = Document(join(dirname(__file__), str(file_name)))
    for paragraph in doc.paragraphs:
        if "SONLI1" in paragraph.text:
            await replace_text(paragraph, "SONLI1", contract_number)
        if "DATE" in paragraph.text:
            await replace_text(paragraph, "DATE", f"{datetime.now().strftime('%d.%m.')}", bold=False)
        if "NAME" in paragraph.text:
            await replace_text(paragraph, "NAME", f"{name}")
        if "FACULTY" in paragraph.text:
            await replace_text(paragraph, "FACULTY", f"{faculty}")
        if "SS!" in paragraph.text:
            await replace_text(paragraph, "SER@", f"{passport}")
        if "ADDRES" in paragraph.text:
            await replace_text(paragraph, "PASSR", f"{address}")
        if "TELNUMBER" in paragraph.text:
            await replace_text(paragraph, "TELNUMBER", f"{number}")
    for table in doc.tables:
        await replace_table_text(table=table, old_text="RFS", new_text=passport, font_size=12)
        await replace_table_text(table=table, old_text="FISH", new_text=name, font_size=12)
        await replace_table_text(table=table, old_text="TELNUMBER", new_text=number, font_size=12)
        await replace_table_text(table=table, old_text="ADDRES", new_text=address, font_size=12)
        await replace_table_text(table=table, old_text="&", new_text=name, bold=False)
    doc.save(join(dirname(__file__), f"file_shartnoma\\{name}.docx"))
    await convert_pdf(name)


async def func_qrcode(url, name, status: bool = False):
    qr = QRCode(
        version=1,
        error_correction=constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )

    qr.add_data(f"https://swine-viable-luckily.ngrok-free.app/get_file/{url}")
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")

    return img.save(join(dirname(__file__), f"file_qrcode/{name}.png" if status else f"ariza_qrcode/{name}.png"))


async def convert_pdf(name, status: bool = False):
    directory = "file_ariza" if status else "file_shartnoma"
    source_path = join(dirname(__file__), f"{directory}\\{name}.docx")
    target_path = join(dirname(__file__), f"{directory}\\{name}.pdf")
    convert(source_path, target_path)
    remove(source_path) if exists(source_path) else None


async def write_qabul(data):
    try:
        path = await get_file_path("file_database\\qabul.xlsx")
        workbook = load_workbook(path)
        sheet = workbook.active
        for row in data:
            sheet.append(row)
        workbook.save(path)
        return True
    except Exception as e:
        return False

# from PyPDF2 import PdfReader, PdfWriter
# import datetime
#
#
# def sign_pdf(input_pdf_path, output_pdf_path, signer_name):
#     # PDF faylini yuklash
#     with open(input_pdf_path, 'rb') as input_pdf_file:
#         reader = PdfReader(input_pdf_file)
#         writer = PdfWriter()
#
#         # Barcha sahifalarni nusxa olish
#         for page in reader.pages:
#             writer.add_page(page)
#
#         # Imzo qismi yaratish
#         signer_block = f"Signed by: {signer_name}\nDate: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
#         # Imzo qismini PDF-ga qo'shish
#         writer.add_text(50, 50, signer_block, fontname='Helvetica', fontsize=12)
#
#         # Imzo qilingan PDF-ni yaratish
#         with open(output_pdf_path, 'wb') as output_pdf_file:
#             writer.write(output_pdf_file)
#
#
# # Imzo qilingan PDF faylini yaratish
# input_pdf_path = 'input.pdf'
# output_pdf_path = 'output_signed.pdf'
# signer_name = 'John Doe'
#
# sign_pdf(input_pdf_path, output_pdf_path, signer_name)
